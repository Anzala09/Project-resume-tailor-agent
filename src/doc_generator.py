"""                                                                                                        
  doc_generator.py — Convert tailored resume text into a formatted .docx file.                                
  """                                                                                                        
                                                                                                             
import re                                                                                                  
from pathlib import Path                                                                                    
   
from docx import Document                                                                                  
from docx.shared import Pt, RGBColor                      
from docx.enum.text import WD_ALIGN_PARAGRAPH
                                                                                                             
                                             
SECTION_HEADERS = {                                                                                        
      "PROFESSIONAL SUMMARY",                              
      "TECHNICAL SKILLS",                                                                                    
      "WORK EXPERIENCE",
      "PROJECTS",                                                                                            
      "EDUCATION",                                          
      "CERTIFICATIONS & ACTIVITIES",
      "CERTIFICATIONS",                      
      "ACTIVITIES",                      
  }
                                                                                                             
   
def _slug(title: str) -> str:                                                                              
      slug = title.lower()                                  
      slug = re.sub(r"[^a-z0-9]+", "_", slug)
      return slug.strip("_")              

                                                                                                             
def _is_section_header(line: str) -> bool:
      return line.strip().upper() in SECTION_HEADERS                                                          
                                                           
                                         
def _build_docx(tailored_text: str) -> Document:
      doc = Document()                                                                                        
      for section in doc.sections:
          section.top_margin = Pt(36)                                                                        
          section.bottom_margin = Pt(36)                    
          section.left_margin = Pt(54)
          section.right_margin = Pt(54)                                                                      
                                         
      lines = tailored_text.split("\n")                                                                      
      in_header_block = True                                                                                  
   
      for i, raw_line in enumerate(lines):                                                                    
          line = raw_line.strip()                          
          if not line:
              if not in_header_block:                                                                        
                  doc.add_paragraph("")
              continue                                                                                        
                                                           
          if i == 0 or (in_header_block and i <= 2 and not _is_section_header(line)):
              p = doc.add_paragraph()                                                                        
              run = p.add_run(line)          
              if i == 0:                                                                                      
                  run.bold = True                          
                  run.font.size = Pt(16)                                                                      
                  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
              else:                                                                                          
                  run.font.size = Pt(10)                    
                  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                  run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)                                            
              continue                        
                                                                                                             
          if _is_section_header(line):                      
              in_header_block = False                                                                        
              p = doc.add_paragraph()
              run = p.add_run(line)                                                                          
              run.bold = True                              
              run.font.size = Pt(11)
              run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
              p.paragraph_format.space_before = Pt(8)
              p.paragraph_format.space_after = Pt(2)
              continue                                                                                        
                                         
          if line.startswith("•") or line.startswith("-") or line.startswith("*"):                            
              in_header_block = False                                                                        
              text = line.lstrip("•-* ").strip()
              p = doc.add_paragraph(style="List Bullet")                                                      
              p.add_run(text).font.size = Pt(10)            
              p.paragraph_format.space_after = Pt(1)
              continue                                                                                        
                                             
          if re.match(r"^[A-Z][A-Za-z\s/,\-\.]+\s*\|", line) or (                                            
              line.isupper() and len(line) < 60            
          ):                                                                                                  
              in_header_block = False        
              p = doc.add_paragraph()                                                                        
              run = p.add_run(line)                        
              run.bold = True                                                                                
              run.font.size = Pt(10)
              p.paragraph_format.space_before = Pt(4)                                                        
              continue                                      
                                         
          if re.search(r"\d{4}", line) and len(line) < 80:
              in_header_block = False                                                                        
              p = doc.add_paragraph()
              run = p.add_run(line)                                                                          
              run.font.size = Pt(9)                        
              run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
              run.italic = True          
              p.paragraph_format.space_after = Pt(1)
              continue                                                                                        
                                         
          in_header_block = False                                                                            
          p = doc.add_paragraph()                                                                            
          run = p.add_run(line)
          run.font.size = Pt(10)                                                                              
          p.paragraph_format.space_after = Pt(2)            
                                             
      return doc                          

                                                                                                             
def generate_document(tailored_text: str, job: dict, output_dir: str) -> str:
      output_path = Path(output_dir)                                                                          
      output_path.mkdir(parents=True, exist_ok=True)        
                                             
      filename = f"resume_alex_morgan_{_slug(job['title'])}"
      docx_path = output_path / f"{filename}.docx"
                                                                                                             
      doc = _build_docx(tailored_text)        
      doc.save(str(docx_path))                                                                                
      print(f"  Saved: {docx_path.name}")                  
                                                                                                             
      pdf_path = output_path / f"{filename}.pdf"
                                                                                                             
      try:                                                  
          from docx2pdf import convert
          convert(str(docx_path), str(pdf_path))
          if pdf_path.exists():              
              print(f"  Converted to PDF: {pdf_path.name}")
              return str(pdf_path)
      except Exception:                                                                                      
          pass
                                                                                                             
      try:                                                  
          import subprocess              
          import shutil
                                                                                                             
          soffice = shutil.which("soffice") or shutil.which("soffice.exe")
          if not soffice:                                                                                    
              win_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
              if Path(win_path).exists():    
                  soffice = win_path      

          if soffice:                                                                                        
              result = subprocess.run(
                  [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_path),                
  str(docx_path)],                                          
                  capture_output=True,    
                  timeout=60,
              )                                                                                              
              if result.returncode == 0 and pdf_path.exists():
                  print(f"  Converted to PDF via LibreOffice: {pdf_path.name}")                              
                  return str(pdf_path)                      
      except Exception:                  
          pass
                                                                                                             
      print(f"  PDF conversion skipped: using .docx instead")
      return str(docx_path)